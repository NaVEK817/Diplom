"""Word report builder for TRACON analysis results."""

from __future__ import annotations

from datetime import datetime
from html import escape
from io import BytesIO
from pathlib import Path
from typing import Iterable
from zipfile import ZIP_DEFLATED, ZipFile


def build_word_report(
    results: dict,
    cluster_rows: list[list],
    output_path: Path,
    generated_at: datetime | None = None,
) -> Path:
    """Create a concise DOCX report for full analysis or algorithm comparison."""
    generated_at = generated_at or datetime.now()
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as error:
        return _build_fallback_docx(results, cluster_rows, output_path, generated_at)

    document = Document()
    _configure_document(document, Pt, Cm)

    if "statistics" in results:
        _build_full_analysis_docx(document, results, cluster_rows, generated_at, output_path)
    else:
        _build_algorithm_docx(document, results, generated_at, output_path)

    for section in document.sections:
        section.start_type = WD_SECTION.CONTINUOUS

    title = document.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(output_path)
    return output_path


def _build_fallback_docx(results: dict, cluster_rows: list[list], output_path: Path, generated_at: datetime) -> Path:
    """Build a valid DOCX with stdlib only when python-docx is unavailable."""
    blocks = []
    images = []
    if "statistics" in results:
        stats = results.get("statistics", {})
        anomaly_stats = stats.get("anomaly_stats", {})
        rule_stats = stats.get("rule_risk_stats", {})
        runway_stats = stats.get("runway_optimization_stats", {})
        what_if = stats.get("what_if_delta", {})
        blocks.append(_p("Выходной документ по анализу траекторий воздушных судов", bold=True, center=True))
        blocks.append(_p(f"Дата формирования: {generated_at.strftime('%Y-%m-%d %H:%M:%S')}"))
        blocks.append(_p("1. Ключевые выводы", bold=True))
        for line in _full_analysis_conclusions(stats, anomaly_stats, rule_stats, runway_stats, what_if):
            blocks.append(_p(f"- {line}"))
        blocks.append(_p("2. Сводные показатели", bold=True))
        blocks.append(_table([["Показатель", "Значение"]] + [
            ["Всего траекторий", stats.get("total_tracks", 0)],
            ["Пучков траекторий", stats.get("total_clusters", 0)],
            ["3D-секторов", stats.get("total_sectors", 0)],
            ["Максимальный индекс нагрузки W", f"{stats.get('max_workload_index', 0):.1f}"],
            ["Конфликтных событий", stats.get("total_conflict_events", 0)],
            ["Аномалий", anomaly_stats.get("total_anomalies_detected", 0)],
            ["Рисков по жестким критериям", rule_stats.get("total_rule_risks", 0)],
            ["PBN/RNP замечаний", stats.get("pbn_issues_count", 0)],
            ["Время выполнения", f"{stats.get('execution_time', 0):.1f} сек"],
        ]))
        chart = _make_pil_full_chart(rule_stats, cluster_rows)
        if chart:
            image_name, image_bytes, width, height = chart
            images.append(("rIdImage1", image_name, image_bytes))
            blocks.append(_p("3. Визуальное представление", bold=True))
            blocks.append(_p("Рисунок 1. Наиболее нагруженные пучки и основные типы рисков."))
            blocks.append(_image("rIdImage1", width, height))
        blocks.append(_p("4. Наиболее значимые пучки", bold=True))
        blocks.append(_table([["Пучок", "Треков", "ВПП", "Аном.", "Риски", "W", "Уровень"]] + _top_cluster_rows(cluster_rows, 10)))
        blocks.append(_p("5. Аномалии и рекомендации", bold=True))
        for item in _top_anomalies(results, limit=8):
            blocks.append(_p(
                f"- Пучок {item['cluster_id']}, трек {item['track_id']}: "
                f"оценка {item['score']:.1%}, риск {item['risk_level']}. {item['description']}"
            ))
    else:
        blocks.append(_p("Выходной документ по сравнению алгоритмов кластеризации", bold=True, center=True))
        blocks.append(_p(f"Дата формирования: {generated_at.strftime('%Y-%m-%d %H:%M:%S')}"))
        rows = [["Алгоритм", "Пучков", "Траекторий", "Средний размер", "Качество", "Время, сек"]]
        for algorithm, result in results.items():
            if isinstance(result, dict) and "error" not in result:
                rows.append([
                    algorithm,
                    result.get("num_clusters", 0),
                    result.get("total_tracks", 0),
                    f"{result.get('avg_size', 0):.1f}",
                    f"{result.get('avg_quality', 0):.1%}",
                    f"{result.get('time', 0):.1f}",
                ])
        blocks.append(_p("1. Сравнительная таблица", bold=True))
        blocks.append(_table(rows))
        chart = _make_pil_algorithm_chart(rows[1:])
        if chart:
            image_name, image_bytes, width, height = chart
            images.append(("rIdImage1", image_name, image_bytes))
            blocks.append(_p("2. Визуальное сравнение", bold=True))
            blocks.append(_image("rIdImage1", width, height))
        blocks.append(_p("3. Вывод", bold=True))
        best_quality = _best_algorithm(results, "avg_quality")
        fastest = _best_algorithm(results, "time", reverse=False)
        if best_quality:
            blocks.append(_p(f"- Лучшее среднее качество кластеров показывает {best_quality}."))
        if fastest:
            blocks.append(_p(f"- Минимальное время выполнения показывает {fastest}."))

    document_xml = _document_xml("".join(blocks))
    _write_docx_package(output_path, document_xml, images)
    return output_path


def _configure_document(document, Pt, Cm):
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(14)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14 if style_name != "Title" else 16)
        style.font.bold = True


def _build_full_analysis_docx(document, results: dict, cluster_rows: list[list], generated_at: datetime, output_path: Path):
    stats = results.get("statistics", {})
    anomaly_stats = stats.get("anomaly_stats", {})
    rule_stats = stats.get("rule_risk_stats", {})
    runway_stats = stats.get("runway_optimization_stats", {})
    what_if = stats.get("what_if_delta", {})

    document.add_heading("Выходной документ по анализу траекторий воздушных судов", 0)
    document.add_paragraph(f"Дата формирования: {generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(
        "Документ содержит сводные результаты полного анализа: секторизацию траекторий, "
        "оценку нагрузки диспетчера, поиск аномалий, PBN/RNP-проверку и рекомендации по дальнейшей обработке."
    )

    document.add_heading("1. Ключевые выводы", level=1)
    for line in _full_analysis_conclusions(stats, anomaly_stats, rule_stats, runway_stats, what_if):
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("2. Сводные показатели", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Показатель"
    table.rows[0].cells[1].text = "Значение"
    _add_metric_rows(
        table,
        [
            ("Всего траекторий", stats.get("total_tracks", 0)),
            ("Пучков траекторий", stats.get("total_clusters", 0)),
            ("3D-секторов", stats.get("total_sectors", 0)),
            ("Максимальный индекс нагрузки W", f"{stats.get('max_workload_index', 0):.1f}"),
            ("Конфликтных событий", stats.get("total_conflict_events", 0)),
            ("Аномалий", anomaly_stats.get("total_anomalies_detected", 0)),
            ("Доля аномалий", f"{anomaly_stats.get('anomaly_percentage', 0):.1f}%"),
            ("Рисков по жестким критериям", rule_stats.get("total_rule_risks", 0)),
            ("PBN/RNP замечаний", stats.get("pbn_issues_count", 0)),
            ("Средняя оценка качества процедуры", f"{stats.get('avg_procedure_quality', 0):.1f}"),
            ("Время выполнения", f"{stats.get('execution_time', 0):.1f} сек"),
        ],
    )

    chart_path = _make_full_analysis_chart(stats, rule_stats, cluster_rows, output_path)
    if chart_path:
        document.add_heading("3. Визуальное представление", level=1)
        document.add_paragraph("Рисунок 1 показывает наиболее нагруженные пучки и распределение основных типов рисков.")
        document.add_picture(str(chart_path), width=_picture_width())
        _safe_unlink(chart_path)

    document.add_heading("4. Наиболее значимые пучки", level=1)
    document.add_paragraph("В таблицу включены пучки с максимальной нагрузкой и/или количеством рисков, чтобы сохранить отчет в пределах 3-4 страниц.")
    top_rows = _top_cluster_rows(cluster_rows, limit=10)
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Пучок", "Треков", "ВПП", "Аном.", "Риски", "W", "Уровень"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in top_rows:
        cells = table.add_row().cells
        for index, value in enumerate(row[:7]):
            cells[index].text = str(value)

    document.add_heading("5. Аномалии и рекомендации", level=1)
    for item in _top_anomalies(results, limit=8):
        document.add_paragraph(
            f"Пучок {item['cluster_id']}, трек {item['track_id']}: "
            f"оценка {item['score']:.1%}, риск {item['risk_level']}. {item['description']}",
            style="List Bullet",
        )

    document.add_paragraph(
        "Рекомендуется приоритизировать проверку пучков с высоким W, большим числом жестких рисков "
        "и повторяющимися PBN/RNP замечаниями; далее уточнить границы секторов и параметры интервалов ВПП."
    )


def _build_algorithm_docx(document, results: dict, generated_at: datetime, output_path: Path):
    document.add_heading("Выходной документ по сравнению алгоритмов кластеризации", 0)
    document.add_paragraph(f"Дата формирования: {generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(
        "Документ сравнивает варианты кластеризации траекторий по числу пучков, охвату, средней величине пучка, качеству и времени выполнения."
    )

    document.add_heading("1. Сравнительная таблица", level=1)
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Алгоритм", "Пучков", "Траекторий", "Средний размер", "Качество", "Время, сек"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    rows = []
    for algorithm, result in results.items():
        if not isinstance(result, dict) or "error" in result:
            continue
        row = [
            algorithm,
            result.get("num_clusters", 0),
            result.get("total_tracks", 0),
            f"{result.get('avg_size', 0):.1f}",
            f"{result.get('avg_quality', 0):.1%}",
            f"{result.get('time', 0):.1f}",
        ]
        rows.append(row)
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    chart_path = _make_algorithm_chart(rows, output_path)
    if chart_path:
        document.add_heading("2. Визуальное сравнение", level=1)
        document.add_picture(str(chart_path), width=_picture_width())
        _safe_unlink(chart_path)

    best_quality = _best_algorithm(results, "avg_quality")
    fastest = _best_algorithm(results, "time", reverse=False)
    document.add_heading("3. Вывод", level=1)
    if best_quality:
        document.add_paragraph(f"Лучшее среднее качество кластеров показывает {best_quality}.", style="List Bullet")
    if fastest:
        document.add_paragraph(f"Минимальное время выполнения показывает {fastest}.", style="List Bullet")
    document.add_paragraph(
        "Для итогового проектного анализа целесообразно выбирать алгоритм не только по времени, "
        "но и по устойчивости количества пучков и интерпретируемости полученных траекторных групп."
    )


def _add_metric_rows(table, rows: Iterable[tuple[str, object]]):
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)


def _picture_width():
    from docx.shared import Cm

    return Cm(15.5)


def _full_analysis_conclusions(stats, anomaly_stats, rule_stats, runway_stats, what_if) -> list[str]:
    overloaded = ", ".join(runway_stats.get("overloaded_runways", [])) or "перегруженные ВПП не выявлены"
    return [
        f"Обработано {stats.get('total_tracks', 0)} траекторий; выделено {stats.get('total_clusters', 0)} пучков и {stats.get('total_sectors', 0)} 3D-секторов.",
        f"Максимальная расчетная нагрузка W = {stats.get('max_workload_index', 0):.1f}; конфликтных событий: {stats.get('total_conflict_events', 0)}.",
        f"Выявлено {anomaly_stats.get('total_anomalies_detected', 0)} аномалий ({anomaly_stats.get('anomaly_percentage', 0):.1f}% от проанализированных траекторий).",
        f"Зафиксировано {rule_stats.get('total_rule_risks', 0)} рисков по жестким критериям на {rule_stats.get('tracks_with_rule_risks', 0)} траекториях.",
        f"По ВПП: прибытий {runway_stats.get('arrivals_count', 0)}, нарушений интервала {runway_stats.get('spacing_violations', 0)}, {overloaded}.",
        f"What-if сценарий изменяет максимальную нагрузку на {what_if.get('max_workload_index', 0):.1f} и число конфликтов на {what_if.get('total_conflicts', 0)}.",
    ]


def _top_cluster_rows(cluster_rows: list[list], limit: int) -> list[list]:
    def score(row):
        workload = _to_float(row[5])
        risks = _to_float(row[4])
        anomalies = _to_float(row[3])
        return workload + risks * 5 + anomalies * 100

    return sorted(cluster_rows, key=score, reverse=True)[:limit]


def _top_anomalies(results: dict, limit: int) -> list[dict]:
    items = []
    for cluster_id, cluster_data in results.get("anomalies", {}).items():
        for item in cluster_data.get("results", []):
            items.append(
                {
                    "cluster_id": cluster_id,
                    "track_id": item.get("track_id", "н/д"),
                    "score": _to_float(item.get("anomaly_score", 0)),
                    "risk_level": item.get("risk_level", "н/д"),
                    "description": item.get("risk_description") or item.get("recommendation") or "Требуется дополнительная проверка.",
                }
            )
    return sorted(items, key=lambda item: item["score"], reverse=True)[:limit]


def _make_full_analysis_chart(stats: dict, rule_stats: dict, cluster_rows: list[list], output_path: Path) -> Path | None:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return None

    top_clusters = _top_cluster_rows(cluster_rows, 6)
    risk_counts = sorted(rule_stats.get("risk_type_counts", {}).items(), key=lambda item: item[1], reverse=True)[:5]
    if not top_clusters and not risk_counts:
        return None

    figure, axes = plt.subplots(1, 2, figsize=(10.5, 4.2))
    if top_clusters:
        axes[0].bar([str(row[0]) for row in top_clusters], [_to_float(row[5]) for row in top_clusters], color="#2f6690")
        axes[0].set_title("Топ пучков по нагрузке W")
        axes[0].set_xlabel("Пучок")
        axes[0].set_ylabel("W")
    else:
        axes[0].axis("off")

    if risk_counts:
        labels = [_short_label(label) for label, _ in risk_counts]
        axes[1].barh(labels, [count for _, count in risk_counts], color="#b23a48")
        axes[1].invert_yaxis()
        axes[1].set_title("Основные типы рисков")
    else:
        axes[1].axis("off")

    figure.tight_layout()
    chart_path = output_path.with_suffix(".summary.png")
    figure.savefig(chart_path, dpi=180, bbox_inches="tight")
    plt.close(figure)
    return chart_path


def _make_algorithm_chart(rows: list[list], output_path: Path) -> Path | None:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return None
    if not rows:
        return None

    algorithms = [str(row[0]) for row in rows]
    qualities = [_to_float(str(row[4]).replace("%", "")) for row in rows]
    times = [_to_float(row[5]) for row in rows]
    figure, axes = plt.subplots(1, 2, figsize=(10.5, 4.2))
    axes[0].bar(algorithms, qualities, color="#357266")
    axes[0].set_title("Среднее качество, %")
    axes[0].tick_params(axis="x", rotation=25)
    axes[1].bar(algorithms, times, color="#c1666b")
    axes[1].set_title("Время выполнения, сек")
    axes[1].tick_params(axis="x", rotation=25)
    figure.tight_layout()
    chart_path = output_path.with_suffix(".summary.png")
    figure.savefig(chart_path, dpi=180, bbox_inches="tight")
    plt.close(figure)
    return chart_path


def _best_algorithm(results: dict, metric: str, reverse: bool = True) -> str | None:
    candidates = [
        (name, result.get(metric, 0))
        for name, result in results.items()
        if isinstance(result, dict) and "error" not in result
    ]
    if not candidates:
        return None
    return sorted(candidates, key=lambda item: item[1], reverse=reverse)[0][0]


def _short_label(value: object, max_len: int = 24) -> str:
    text = str(value)
    return text if len(text) <= max_len else text[: max_len - 1] + "."


def _to_float(value: object) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def _safe_unlink(path: Path):
    try:
        path.unlink(missing_ok=True)
    except OSError:
        pass


def _make_pil_full_chart(rule_stats: dict, cluster_rows: list[list]):
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None

    top_clusters = _top_cluster_rows(cluster_rows, 6)
    risk_counts = sorted(rule_stats.get("risk_type_counts", {}).items(), key=lambda item: item[1], reverse=True)[:5]
    if not top_clusters and not risk_counts:
        return None

    width, height = 1200, 520
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = _pil_font(ImageFont, 22)
    small = _pil_font(ImageFont, 18)
    draw.text((36, 20), "Топ пучков по нагрузке W", fill="#1f2933", font=font)
    draw.text((650, 20), "Основные типы рисков", fill="#1f2933", font=font)

    if top_clusters:
        max_w = max(_to_float(row[5]) for row in top_clusters) or 1
        bar_w = 62
        base_y = 430
        for index, row in enumerate(top_clusters):
            x = 60 + index * 82
            bar_h = int((_to_float(row[5]) / max_w) * 310)
            draw.rectangle((x, base_y - bar_h, x + bar_w, base_y), fill="#2f6690")
            draw.text((x, base_y + 12), str(row[0]), fill="#111827", font=small)
        draw.line((44, base_y, 580, base_y), fill="#6b7280", width=2)

    if risk_counts:
        max_count = max(count for _, count in risk_counts) or 1
        for index, (label, count) in enumerate(risk_counts):
            y = 78 + index * 70
            bar_w = int((count / max_count) * 360)
            draw.rectangle((770, y, 770 + bar_w, y + 28), fill="#b23a48")
            draw.text((650, y), _short_label(label, 18), fill="#111827", font=small)
            draw.text((780 + bar_w, y), str(count), fill="#111827", font=small)

    return _pil_image_payload(image, "summary.png")


def _make_pil_algorithm_chart(rows: list[list]):
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None
    if not rows:
        return None

    width, height = 1200, 500
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = _pil_font(ImageFont, 22)
    small = _pil_font(ImageFont, 18)
    draw.text((36, 20), "Среднее качество, %", fill="#1f2933", font=font)
    draw.text((650, 20), "Время выполнения, сек", fill="#1f2933", font=font)

    qualities = [_to_float(str(row[4]).replace("%", "")) for row in rows]
    times = [_to_float(row[5]) for row in rows]
    max_quality = max(qualities) or 1
    max_time = max(times) or 1
    base_y = 410
    for index, row in enumerate(rows):
        x = 70 + index * 92
        q_h = int((qualities[index] / max_quality) * 300)
        draw.rectangle((x, base_y - q_h, x + 54, base_y), fill="#357266")
        draw.text((x - 12, base_y + 12), _short_label(row[0], 8), fill="#111827", font=small)

        tx = 690 + index * 92
        t_h = int((times[index] / max_time) * 300)
        draw.rectangle((tx, base_y - t_h, tx + 54, base_y), fill="#c1666b")
        draw.text((tx - 12, base_y + 12), _short_label(row[0], 8), fill="#111827", font=small)
    draw.line((44, base_y, 565, base_y), fill="#6b7280", width=2)
    draw.line((625, base_y, 1155, base_y), fill="#6b7280", width=2)
    return _pil_image_payload(image, "summary.png")


def _pil_font(ImageFont, size: int):
    for name in ("times.ttf", "arial.ttf"):
        try:
            return ImageFont.truetype(name, size)
        except Exception:
            pass
    return ImageFont.load_default()


def _pil_image_payload(image, name: str):
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return name, buffer.getvalue(), image.width, image.height


def _p(text: object, bold: bool = False, center: bool = False) -> str:
    justification = '<w:jc w:val="center"/>' if center else ""
    bold_tag = "<w:b/>" if bold else ""
    return (
        "<w:p>"
        f"<w:pPr>{justification}</w:pPr>"
        "<w:r>"
        f"<w:rPr><w:rFonts w:ascii=\"Times New Roman\" w:hAnsi=\"Times New Roman\" w:cs=\"Times New Roman\"/><w:sz w:val=\"28\"/>{bold_tag}</w:rPr>"
        f"<w:t xml:space=\"preserve\">{escape(str(text))}</w:t>"
        "</w:r>"
        "</w:p>"
    )


def _table(rows: list[list]) -> str:
    cells_xml = []
    for row in rows:
        row_xml = ["<w:tr>"]
        for value in row:
            row_xml.append(
                "<w:tc><w:tcPr><w:tcW w:w=\"2400\" w:type=\"dxa\"/></w:tcPr>"
                f"{_p(value)}"
                "</w:tc>"
            )
        row_xml.append("</w:tr>")
        cells_xml.append("".join(row_xml))
    return (
        "<w:tbl>"
        "<w:tblPr><w:tblW w:w=\"0\" w:type=\"auto\"/>"
        "<w:tblBorders><w:top w:val=\"single\" w:sz=\"4\"/><w:left w:val=\"single\" w:sz=\"4\"/>"
        "<w:bottom w:val=\"single\" w:sz=\"4\"/><w:right w:val=\"single\" w:sz=\"4\"/>"
        "<w:insideH w:val=\"single\" w:sz=\"4\"/><w:insideV w:val=\"single\" w:sz=\"4\"/></w:tblBorders>"
        "</w:tblPr>"
        + "".join(cells_xml)
        + "</w:tbl>"
    )


def _image(rel_id: str, width_px: int, height_px: int) -> str:
    width_emu = 5580000
    height_emu = max(1, int(width_emu * height_px / max(1, width_px)))
    return f"""
<w:p>
  <w:r>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:docPr id="1" name="Report chart"/>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic>
              <pic:nvPicPr><pic:cNvPr id="0" name="chart.png"/><pic:cNvPicPr/></pic:nvPicPr>
              <pic:blipFill><a:blip r:embed="{rel_id}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>
              <pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>
"""


def _document_xml(body: str) -> str:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
  xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
  <w:body>
    {body}
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="907" w:right="1021" w:bottom="907" w:left="1021" w:header="708" w:footer="708" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>
"""


def _write_docx_package(output_path: Path, document_xml: str, images: list[tuple[str, str, bytes]] | None = None):
    images = images or []
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Default Extension="png" ContentType="image/png"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
    image_rels = "".join(
        f'<Relationship Id="{rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/{name}"/>'
        for rel_id, name, _ in images
    )
    document_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{image_rels}</Relationships>
"""
    styles = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:rPr>
      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
      <w:sz w:val="28"/>
    </w:rPr>
  </w:style>
</w:styles>
"""
    with ZipFile(output_path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels)
        archive.writestr("word/_rels/document.xml.rels", document_rels)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/styles.xml", styles)
        for _, name, image_bytes in images:
            archive.writestr(f"word/media/{name}", image_bytes)
